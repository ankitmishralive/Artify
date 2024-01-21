from PyPDF2 import PdfReader
from deep_translator import GoogleTranslator
import docx
from docx2pdf import convert
import pythoncom

pdf_file = "Ankit Mishra.pdf"

# Read PDF
reader = PdfReader(pdf_file)

txt = ""

# Extract text from PDF
for page in reader.pages:
    text = page.extract_text()
    if text:
        txt += text

from  textblob import TextBlob as TB

_text = TB(txt)
iso_code = _text.detect_language()
print(f'Language Detected:{iso_code}')
# Translate to Hindi
# Translated_text = GoogleTranslator(source='auto', target='hi').translate(txt) 
# # Translated_text = Translated_text.text

# # Create a Word document and save the translated text
# doc = docx.Document()
# # doc.add_heading('Translated Text', 3)
# doc.add_paragraph(Translated_text, style='BodyText')

# # Save the document
# doc.save('translated_text.docx')



# convert('translated_text.docx','static/Translater/Translated.pdf',pythoncom.CoInitialize())


#-----------------------For Knowing Languages--------------- 


# import googletrans


# print(googletrans.LANGUAGES)
